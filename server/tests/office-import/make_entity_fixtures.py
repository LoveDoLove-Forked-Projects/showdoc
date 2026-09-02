#!/usr/bin/env python3
"""
生成 HTML 实体多重编码攻击样本（供 http_entity_test.php 使用）。

  entity_attack.docx   正文含双重/三重实体编码的 <script>/<img onerror>/javascript:
  entity_normal.docx   正常功能对照文档（H1 + 2×H2 拆页）
"""
import os

import docx

out = os.environ.get("FX_OUT") or os.path.join(os.path.dirname(os.path.abspath(__file__)), "fixtures")
os.makedirs(out, exist_ok=True)

# 多重实体编码载荷（docx 正文中以纯文本形态写入，转换后须被 sanitizer 完全中和）
doc = docx.Document()
doc.add_heading("Entity Verify", level=1)
for line in [
    '&amp;lt;script&amp;gt;alert("double-script")&amp;lt;/script&amp;gt;',
    '&amp;lt;img src=x onerror=alert("double-img")&amp;gt;',
    '[double](&amp;#106;avascript:alert("double-js"))',
    '&amp;amp;lt;script&amp;amp;gt;alert("triple-script")&amp;amp;lt;/script&amp;amp;gt;',
    '&amp;amp;lt;img src=x onerror=alert("triple-img")&amp;amp;gt;',
    '[triple](&amp;amp;#106;avascript:alert("triple-js"))',
    '正常段落，含 [正常链接](https://showdoc.com.cn) 与 **加粗**。',
]:
    doc.add_paragraph(line)
doc.save(os.path.join(out, "entity_attack.docx"))
print("entity_attack.docx")

# 正常功能对照
doc2 = docx.Document()
doc2.add_heading("Normal Doc", level=1)
doc2.add_paragraph("这是正常功能验证文档。")
doc2.add_heading("Section A", level=2)
doc2.add_paragraph("Section A 内容，含 [链接](https://example.com) 和 `code`。")
doc2.add_heading("Section B", level=2)
doc2.add_paragraph("Section B 内容。")
doc2.save(os.path.join(out, "entity_normal.docx"))
print("entity_normal.docx")
